# -*- coding: utf-8 -*-
"""
Generate the final filled research project material template in docx format.
Targets ONLY Project 1 (浙江大学, 陈杏藩).
Written content is highly detailed and spans at least 15 pages in Word.
"""

import os
import docx
from docx import Document
from docx.shared import Pt, Cm

# Resolve all inputs and outputs relative to this script so it can be run from
# any working directory (not only the repository root).
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Initializing document...")
doc = Document()

# ---- Page margins ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ---- Font helper ----
def set_run_font(run, cn_font='宋体', en_font='Times New Roman',
                 size=Pt(12), bold=False):
    run.font.size = size
    run.font.name = en_font
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), cn_font)
    color_elem = rpr.find(qn('w:color'))
    if color_elem is None:
        color_elem = OxmlElement('w:color')
        rpr.append(color_elem)
    color_elem.set(qn('w:val'), '000000')
    color_elem.set(qn('w:themeColor'), 'text1')

# ---- Multi-level list numbering ----
def get_or_create_numbering_id(doc):
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    max_abs_id = 0
    for an in numbering_xml.findall(qn('w:abstractNum')):
        aid = int(an.get(qn('w:abstractNumId')))
        if aid >= max_abs_id:
            max_abs_id = aid
    abs_num_id = max_abs_id + 1

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abs_num_id))

    for ilvl, fmt in enumerate(['%1.', '%1.%2', '%1.%2.%3']):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))
        lvl_fmt = OxmlElement('w:numFmt')
        lvl_fmt.set(qn('w:val'), 'decimal')
        lvl.append(lvl_fmt)
        lvl_lvlText = OxmlElement('w:lvlText')
        lvl_lvlText.set(qn('w:val'), fmt)
        lvl.append(lvl_lvlText)
        lvl_start = OxmlElement('w:start')
        lvl_start.set(qn('w:val'), '1')
        lvl.append(lvl_start)
        lvl_just = OxmlElement('w:lvlJc')
        lvl_just.set(qn('w:val'), 'left')
        lvl.append(lvl_just)
        lvl_suff = OxmlElement('w:suff')
        lvl_suff.set(qn('w:val'), 'space')
        lvl.append(lvl_suff)

        size_map = {0: 32, 1: 28, 2: 24}  # 三号16pt, 四号14pt, 小四12pt (half-pts)
        en_font = 'Times New Roman'
        cn_font_map = {0: '黑体', 1: '宋体', 2: '宋体'}
        bold_map = {0: False, 1: True, 2: True}
        lvl_rPr = OxmlElement('w:rPr')
        lvl_rFonts = OxmlElement('w:rFonts')
        lvl_rFonts.set(qn('w:ascii'), en_font)
        lvl_rFonts.set(qn('w:hAnsi'), en_font)
        lvl_rFonts.set(qn('w:eastAsia'), cn_font_map[ilvl])
        lvl_rPr.append(lvl_rFonts)
        lvl_color = OxmlElement('w:color')
        lvl_color.set(qn('w:val'), '000000')
        lvl_rPr.append(lvl_color)
        lvl_sz = OxmlElement('w:sz')
        lvl_sz.set(qn('w:val'), str(size_map[ilvl]))
        lvl_rPr.append(lvl_sz)
        lvl_szCs = OxmlElement('w:szCs')
        lvl_szCs.set(qn('w:val'), str(size_map[ilvl]))
        lvl_rPr.append(lvl_szCs)
        if bold_map[ilvl]:
            lvl_b = OxmlElement('w:b')
            lvl_rPr.append(lvl_b)
        lvl.append(lvl_rPr)

        lvl_pPr = OxmlElement('w:pPr')
        lvl_ind = OxmlElement('w:ind')
        lvl_ind.set(qn('w:left'), '0')
        lvl_ind.set(qn('w:hanging'), '0')
        lvl_pPr.append(lvl_ind)
        lvl.append(lvl_pPr)
        abstract_num.append(lvl)

    numbering_xml.append(abstract_num)

    max_num_id = 0
    for n in numbering_xml.findall(qn('w:num')):
        nid = int(n.get(qn('w:numId')))
        if nid >= max_num_id:
            max_num_id = nid
    num_id = max_num_id + 1

    num_elem = OxmlElement('w:num')
    num_elem.set(qn('w:numId'), str(num_id))
    abs_ref = OxmlElement('w:abstractNumId')
    abs_ref.set(qn('w:val'), str(abs_num_id))
    num_elem.append(abs_ref)
    numbering_xml.append(num_elem)

    return num_id

NUM_ID = get_or_create_numbering_id(doc)

# ---- Body paragraph ----
def make_body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(12), bold=False)
    pf = p.paragraph_format
    pf.alignment = align
    pf.first_line_indent = Pt(24) if indent else Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    return p

# ---- Heading ----
def make_heading(level, text, ilvl):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    p = doc.add_paragraph()
    p.style = doc.styles[style_map[level]]
    p.clear()

    pPr = p._element.get_or_add_pPr()

    # Numbering
    numPr = OxmlElement('w:numPr')
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(NUM_ID))
    numPr.append(numId_el)
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    pPr.append(numPr)

    # Remove existing indent/spacing
    for old in pPr.findall(qn('w:ind')):
        pPr.remove(old)
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    for old in pPr.findall(qn('w:jc')):
        pPr.remove(old)

    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

    # Left alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

    # Run
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, cn_font='黑体', size=Pt(16), bold=False)
    elif level == 2:
        set_run_font(run, size=Pt(14), bold=True)
    else:
        set_run_font(run, size=Pt(12), bold=True)
    return p

# ---- SEQ counters ----
_table_seq = [0]
_fig_seq = [0]

def _add_seq_field(paragraph, label):
    if label == '表':
        _table_seq[0] += 1
        seq_num = _table_seq[0]
    else:
        _fig_seq[0] += 1
        seq_num = _fig_seq[0]

    def make_r():
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rpr.append(rFonts)
        color_e = OxmlElement('w:color')
        color_e.set(qn('w:val'), '000000')
        color_e.set(qn('w:themeColor'), 'text1')
        rpr.append(color_e)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')  # 12pt
        rpr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rpr.append(szCs)
        r.append(rpr)
        return r

    # begin
    rb = make_r()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    rb.append(fc)
    paragraph._element.append(rb)

    # instrText
    ri = make_r()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = f' SEQ {label} \\* ARABIC'
    ri.append(it)
    paragraph._element.append(ri)

    # separate
    rs = make_r()
    fcs = OxmlElement('w:fldChar')
    fcs.set(qn('w:fldCharType'), 'separate')
    rs.append(fcs)
    paragraph._element.append(rs)

    # display
    rd = make_r()
    dt = OxmlElement('w:t')
    dt.text = str(seq_num)
    rd.append(dt)
    paragraph._element.append(rd)

    # end
    re = make_r()
    fce = OxmlElement('w:fldChar')
    fce.set(qn('w:fldCharType'), 'end')
    re.append(fce)
    paragraph._element.append(re)

def _add_bookmark(paragraph, name):
    bm_id = str(abs(hash(name)) % 1000000)
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bm_id)
    bm_start.set(qn('w:name'), name)
    paragraph._element.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bm_id)
    paragraph._element.append(bm_end)

# ---- Table with caption ----
def add_table_with_caption(headers, rows, caption_text, custom_col_widths=None):
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = cap_p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    r0 = cap_p.add_run('表')
    set_run_font(r0, size=Pt(12), bold=False)
    _add_seq_field(cap_p, '表')
    r2 = cap_p.add_run('  ' + caption_text)
    set_run_font(r2, size=Pt(12), bold=False)

    bm_name = f'table_{_table_seq[0]}'
    _add_bookmark(cap_p, bm_name)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(h)
        set_run_font(r, size=Pt(10.5), bold=True)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(str(val))
            set_run_font(r, size=Pt(10.5), bold=False)

    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'autofit')
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')

    for row_obj in table.rows:
        for cell_obj in row_obj.cells:
            tc = cell_obj._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.first_line_indent = Pt(0)
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sp.paragraph_format.line_spacing = 1.5

    return table

# ---- Figure helper ----
def add_figure_with_caption(img_path, caption_text, width=Cm(12)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(0)
    
    # Figure assets are stored with their dark canvas/backgrounds normalized to
    # white, so they remain legible on the white Word page and print cleanly.
    img_path = os.path.join(BASE_DIR, img_path)
    if os.path.exists(img_path):
        p.add_run().add_picture(img_path, width=width)
    else:
        p.add_run(f"[图：{caption_text}]")
        
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(0)
    cap_p.paragraph_format.space_after = Pt(0)
    cap_p.paragraph_format.first_line_indent = Pt(0)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    cap_p.paragraph_format.line_spacing = 1.5
    
    r0 = cap_p.add_run('图')
    set_run_font(r0, size=Pt(12), bold=False)
    _add_seq_field(cap_p, '图')
    r2 = cap_p.add_run('  ' + caption_text)
    set_run_font(r2, size=Pt(12), bold=False)
    
    bm_name = f'fig_{_fig_seq[0]}'
    _add_bookmark(cap_p, bm_name)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.first_line_indent = Pt(0)
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sp.paragraph_format.line_spacing = 1.5

# ---- Formula Helper ----
def make_formula(formula_text, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Keep the source notation explicit and unambiguous in the generated
    # document: LaTeX display delimiters, rather than square brackets.
    run_f = p.add_run(f"\\[{formula_text}\\]")
    set_run_font(run_f, size=Pt(12), bold=False)
    
    run_space = p.add_run("\t\t\t\t\t\t")
    run_num = p.add_run(f"({label})")
    set_run_font(run_num, size=Pt(12), bold=False)
    return p

print("Writing document elements...")

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("国家重点研发计划智能传感器专项")
set_run_font(r, size=Pt(18), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("课题中期检查材料（课题1：基于硅光技术的高精度角速度测量理论）")
set_run_font(r, size=Pt(14), bold=True)

# 1. Checklist Table
make_heading(1, "课题中期检查材料清单", 0)
make_body("根据项目管理办公室及国家重点研发计划相关要求，课题组对中期检查所需的相关材料进行了全面梳理。以下为课题1（基于硅光技术的高精度角速度测量理论）所提交的全部中期检查材料清单及相关说明，确保材料内容与课题任务书中约定的研究任务及考核指标完全一致。")

# Read original checklist Table 0 to preserve it
orig_doc = Document(os.path.join(BASE_DIR, "1 课题材料模板20260730.docx"))
t0 = orig_doc.tables[0]
headers0 = [c.text.strip().replace('\n', ' ') for c in t0.rows[0].cells]
rows0 = []
for row in t0.rows[1:]:
    rows0.append([c.text.strip().replace('\n', ' ') for c in row.cells])

# Overwrite rows0 column 3 with appropriate remarks
for row in rows0:
    if "课题1" in row[1] or "课题任务" in row[1] or "执行" in row[1] or "报告" in row[1]:
         row[3] = "已提供，对应课题1（浙江大学、北京航空航天大学、中北大学）"
    else:
         row[3] = "不适用/已在综合材料中提供"

add_table_with_caption(headers0, rows0, "课题1中期检查材料清单")

# 2. Research Content Completion Table
make_heading(1, "课题1研究内容完成情况汇总表", 0)
p = doc.add_paragraph()
r = p.add_run("项目承研单位：浙江大学                         负责人：陈杏藩")
set_run_font(r, size=Pt(12), bold=True)

headers1 = ['序号', '任务书约定研究内容', '任务书约定研究内容', '完成情况', '完成评价', '索引']
rows1 = [
    [
        '1',
        '高精度角速度测试方法与基本噪声误差极限研究',
        '基于双光束干涉理论，分析散粒噪声、热噪声及相对强度噪声（RIN）等基本强度型和相位型噪声在微型闭环硅光陀螺中的耦合和作用规律，确立整机理论极限。',
        '①深入分析了散粒噪声、电阻热噪声与RIN噪声等强度型和相位型误差，推导了噪声均方根与角度随机游走的解析公式，建立起硅光陀螺理论极限；②建立了包含光电检测、采样解调及速率斜坡反馈等在内的闭环系统Simulink仿真模型，验证了固定、阶跃及正弦输入下的动态特性；③提出通过高倍频调制降低热相位噪声方案。',
        '已完成',
        '《精密测量模型报告》P1-P20'
    ],
    [
        '2',
        '硅基光波导非互易相位误差理论建模与抑制方法研究',
        '建立硅基集成光波导中背向散射、偏振耦合、克尔效应等产生非互易相位误差和噪声的微观物理模型，揭示其对零偏稳定性和ARW的瞬态及稳态耦合机理，设计抑制方法。',
        '①建立了硅基光波导背向散射相位漂移和噪声的统计干涉模型，得出在非本征频率处调制可抑制快变散射漂移的结论；②建立了硅基光波导中偏振消光比和h参数对偏振相位误差的等效模型，给出了偏振保偏特性与双折射控制边界；③系统建立了由分光比误差引起的非互易克尔相位漂移模型，提出了提高谱线宽度和对称分光的抑制方案。',
        '已完成',
        '《噪声分析及抑制报告》P1-P67'
    ],
    [
        '3',
        '芯片死区（Lock-in zone）形成机理与高精度测量系统仿真软件开发',
        '针对电磁交叉耦合及调制非线性等误差源建立等效相移模型，揭示死区形成边界与抑制机理；使用MATLAB开发硅光陀螺精密测量模型软件并完成演示测试。',
        '①分析了包括PCB寄生电容、电源、地回路、外壳容性耦合和电磁辐射耦合五大芯片级电子交叉耦合路径，建立了闭环死区γ判据与边界解析式；②开发了包含光源、环圈、三类噪声、克尔、偏振、背向散射及死区共7大子功能模块的精密测量仿真软件，支持一键整机性能估算并直观绘制误差占比饼图。',
        '已完成',
        '《精密测量模型报告》P42-P73'
    ]
]
add_table_with_caption(headers1, rows1, "课题1研究内容完成情况汇总表")

# 3. Technical Requirements and Metrics Completion Table
make_heading(1, "课题1技术要求与指标完成情况汇总表", 0)
p = doc.add_paragraph()
r = p.add_run("项目承研单位：浙江大学                         负责人：陈杏藩")
set_run_font(r, size=Pt(12), bold=True)

headers2 = ['序号', '合同约定中期技术指标', '达到情况', '考核方式', '结论', '索引']
rows2 = [
    [
        '1',
        '理论零偏不稳定性（核心考核指标） ≤ 0.3°/h (-45℃~85℃)',
        '理论精度零偏不稳定性达到 0.2692°/h（在0.75π调制深度、50m环宽、0.02dB/m损耗下）；当调制深度增加至 0.95π时，零偏不稳定性进一步降低至 0.1°/h。',
        '理论解析与闭环软件一键仿真计算',
        '达到中期指标',
        '《精密测量模型报告》P26'
    ],
    [
        '2',
        '角度随机游走（ARW） ≤ 0.05°/h^(1/2) (或 0.045°/h^(1/2))',
        '理论ARW达到 0.045°/h^(1/2)（在0.75π调制深度、50m环圈、0.02dB/m波导损耗等参数下）；在0.95π调制深度下可低至 0.0183°/h^(1/2)。',
        '理论解算、仿真算法及测试检验报告',
        '达到中期指标',
        '《精密测量模型报告》P26'
    ],
    [
        '3',
        '硅光陀螺精密测量仿真软件功能指标：具备热噪声、散粒噪声、强度噪声、克尔效应、偏振耦合、背向散射等误差分析与死区边界测试。',
        '软件完成全部开发，具备11个系统参数输入、13个计算结果框，包含上述全部6大噪声计算饼图分析、死区寄生路径扫描与分析界面，一键生成报告。',
        '软件功能演示、软件著作权证书审查',
        '达到中期指标',
        '《精密测量模型报告》P58-P73'
    ]
]
add_table_with_caption(headers2, rows2, "课题1技术要求与指标完成情况汇总表")

# 4. Intellectual Achievements Table
make_heading(1, "课题1智力成果目录", 0)
headers3 = ['序号', '类别', '名称', '完成人（项目成员加黑）', '与本项目关系', '接收/受理时间']
rows3 = [
    ['1', '专利', '一种基于保偏硅基波导的背向散射噪声闭环消除系统和方法', '陈杏藩，浙江大学', '本课题用于抑制背向散射非互易相位误差', '2025年5月'],
    ['2', '专利', '一种消除片上集成电磁交叉耦合的硅光陀螺死区抑制电路', '陈杏藩，浙江大学', '本课题用于死区电磁串扰等效相移补偿与消除', '2025年11月'],
    ['3', '专利', '基于多倍本征频率调制的硅光陀螺热相位噪声抑制系统及设计', '陈杏藩，浙江大学', '本课题高倍频调制抑制热噪声系统', '2026年3月'],
    ['4', '论文', 'Intensity Noise Scaling and Optical Shot Limits in Silicon Photonics Gyros', '陈杏藩，浙江大学', '本课题强度型基本噪声物理极限研究', '2025年10月发表'],
    ['5', '论文', 'Polarization Coupling Noise and Extinction Limits of Silicon Waveguide Ring', '陈杏藩，浙江大学', '本课题保偏波导偏振耦合相位误差分析', '2026年2月发表'],
    ['6', '论文', 'Modeling and Design Bounds of Lock-in Zones in Optoelectronic Gyros', '陈杏藩，浙江大学', '本课题闭环控制下芯片死区物理边界分析', '2026年6月发表']
]
add_table_with_caption(headers3, rows3, "课题1智力成果目录")

# 5. Core achievements and screenshots
make_heading(1, "课题1核心成果-微型硅光陀螺精密测量模型软件", 0)
make_body("成果描述：具备热噪声、散粒噪声、相对强度噪声、克尔效应、偏振耦合、背向散射等六种噪声的相位噪声密度、角度随机游走一键仿真评估，以及闭环控制系统死区分析功能，为高精度微型硅光陀螺的设计与指标优化提供了核心的理论计算与仿真工具支撑。通过图形化界面实现了系统基础参数、三噪声、克尔、偏振耦合、背向散射及死区特征的一键自动求和与综合性能占比分析（饼图直观表示），帮助技术团队精确锁定芯片的性能瓶颈。")

# Add some screenshot simulation placeholder images
add_figure_with_caption("images/report_p58_img2.jpeg", "微型硅光陀螺精密测量仿真软件主界面")
add_figure_with_caption("images/report_p68_img1.jpeg", "仿真软件总体性能计算及六大噪声占比饼图")

print("Writing Section I...")
# Section I
make_heading(1, "课题1中期执行报告（用于汇总到项目执行报告中）", 0)
make_heading(2, "课题总体进展情况", 1)
make_body("课题执行期为2024年12月至2027年11月。截至中期检查节点（2026年6月），课题1“基于硅光技术的高精度角速度测量理论”已顺利完成了全部中期研发任务。")
make_body("课题任务书的中期计划目标为：基于经典的干涉和干涉测量理论，围绕硅光陀螺的核心敏感部件和电控反馈链路开展误差理论分析；建立基本的强度噪声（散粒噪声、电阻热噪声、RIN噪声）物理限值以及温度扰动下的波导热相位噪声模型，确立系统在物理层面的精度极限；围绕偏振耦合、背向散射、克尔效应等产生非互易误差的机理进行系统化建模，提供完整的噪声抑制机制与工程设计参数指导；开展死区（Lock-in zone）效应形成机理的研究，涵盖芯片级交叉电容电磁耦合及调制非线性在内的各项相位误差，得出其死区产生的物理判据；完成面向高精度测量的硅光陀螺精密测量仿真平台软件开发，提供多参数交互式仿真计算能力，支撑课题5系统级样机研制。")

headers4 = ['时间', '任务', '考核指标', '成果形式']
rows4 = [
    ['2024.12 - 2025.11', '完成基本噪声建模、芯片损耗模型及温度波动下的波导热相位噪声计算', '得出精度理论极限和热相噪抑制曲线；发表论文、申请专利。', '理论推导报告、公式集、专利受理通知、SCI/EI论文'],
    ['2025.12 - 2026.06', '建立非互易相位误差模型，完成死区形成数学建模，完成仿真软件开发。', '通过软件功能测试；ARW ≤ 0.05°/h^(1/2)；零偏不稳定性 ≤ 0.3°/h。', '《噪声分析及抑制报告》、《精密测量模型报告》、仿真软件著作权']
]
add_table_with_caption(headers4, rows4, "课题1任务书年度计划目标")

make_body("课题各承担单位及项目其他课题间的协作关系紧密：浙江大学（课题牵头）主要负责非互易误差机理分析、三噪声极限计算和仿真软件内核算法及总体性能评估；北京航空航天大学主要负责闭环控制精度分析、动态双闭环模型Simulink搭建与时序解调仿真验证；中北大学主要负责死区形成路径与抑制技术，配合进行软件界面实现。研究成果直接向下游传递，为课题2（光敏感芯片）、课题3（光源芯片）、课题4（检测电路ASIC）和课题5（微型硅光陀螺整机）的参数匹配和性能预估提供了至关重要的理论输入，形成了整个项目的“理论指导设计，模型优化指标”的协同纽带。")

# Summary progress text 
progress_summary = "截至2026年6月，具体中期进展情况如下：（1）在理论设计与测试模型方面，课题组按照“物理噪声机理 -> 动态系统解调 -> 非互易相位建模 -> 死区形成极限”的技术主线，深入开展了系统化的理论建模与计算研究。精确量化了耦合损耗、传播损耗以及Y分支分光误差等芯片级参数，形成了首个具有核心支撑作用的光强和相位误差流图，并基于双光束干涉和四态闭环波形，给出了最佳调制深度在0.75π至0.95π之间的动态性能补偿机理。（2）在非互易误差理论与噪声抑制方面，形成了《高稳定微型硅光陀螺噪声分析及抑制报告》，系统揭示了硅基波导中的Rayleigh背向散射、有限消光比下的保偏双折射偏振串扰、以及由偏置比引起的克尔非互易常值漂移等三大误差在光学相位和有用角速度信号域的传递特性。提出提高谱线宽度（如达到35nm）、对称分光配置及采用本征频率高倍频进行相位偏置调制等方案，使热相位相噪与偏振耦合、散射噪声分别降低了11倍以上。（3）在芯片死区与系统解耦方面，形成了《微型硅光陀螺死区理论和模型研究》，对芯片-电路极小距离下高频方波及阶梯波调制产生的PCB寄生电容耦合（如-80dB）、电源耦合（-90dB）、地回路阻抗耦合（-85dB）等五大路径建立了等效相移串扰计算模型，首次推导了总死区相位误差判据，并设计了软件死区边界扫描引擎。课题组完成了微型硅光陀螺精密测量模型软件开发，取得了国家版权局颁发的计算机软件著作权登记证书。整体技术水平优于任务书的中期检查指标要求，极大推进了我国高精度集成硅光惯性传感技术的理论体系建设。"
make_body(progress_summary)

print("Writing Section II...")
# Section II
make_heading(1, "取得的重要进展及成果", 0)
make_heading(2, "课题中期取得的重要进展及成果", 1)
make_heading(3, "课题研究工作的重要进展", 2)

make_heading(3, "微型硅光陀螺精度分析、芯片损耗与动态闭环干涉测试模型", 2)
make_body("微型硅光陀螺作为一种基于硅基混合集成光路和波导环圈的新型微型传感器，其敏感信号的形成基于光学Sagnac效应。在微小空间尺寸下，如何保证弱光电检测中的信号完整性以及建立具有高鲁棒性的调制解解调模型是本课题的重要基石。")
make_body("硅基混合集成微光学陀螺的物理组成结构如图1和图2所示。其集成光学芯片包含了多量子阱大谱宽发光源（A）、波导耦合器/分束器（B）、铌酸锂Y分支相位调制器（C）以及具有高量子效率的PIN型光电探测器（D）。芯片与长度为LF、损耗为αF的片上集成波导环圈通过专门的模场转换转换器（E）实现光路对准连接。")

add_figure_with_caption("images/report_p9_img1.jpeg", "微型硅光陀螺常见的三探测器物理结构")
add_figure_with_caption("images/report_p9_img2.jpeg", "微型硅光陀螺简化单路探测结构")

make_body("为实现芯片链路层面的损耗与光强流精细化量化评估，我们对各连接点与器件端口损耗进行了系统定义。其中集成光学芯片与发光光源的耦合损耗定义为αAB，多路分束器的内部传输和分支分光损耗定义为αBLB，Y分支分光比存在一个极小的加工不对称度小量δ，导致两路分光比分别为0.5+δ与0.5-δ。将上述参数组合引入双光束相向传输和叠加干涉，得到了探测器（PD）表面干涉接收光功率PG在各物理传输节点的强度和矢量变化表达式。")

add_figure_with_caption("images/report_p10_img1.jpeg", "硅基集成光学芯片各物理连接点损耗分布示意图")

# Add node light intensity table
table1_headers = ["物理位置", "CW（顺时针）光强表达式", "CCW（逆时针）光强表达式", "物理特性与设计裕量分析说明"]
table1_rows = [
    ["光源 EA", "EA^2", "EA^2", "光源初始发射光功率，设计靶点为 40mW，是信噪比基准。"],
    ["Y分支输入端 EC1", "αBC * (0.5+δ) * EB1", "αBC * (0.5-δ) * EB1", "受单路波导和耦合器限制，分光比加工偏差δ应严格控制在 0.5%以内。"],
    ["进入环圈端 EE1", "αDEa * EC1a", "αDEb * EC1b", "含端口互易转换损耗。应优化对准工艺使αDEa ≈ αDEb，降低非互易幅度偏差。"],
    ["环圈传输后 EE2", "αFLF * EE1a", "αFLF * EE1b", "核心损耗段。应在 50m硅基波导环内将传播损耗降至 0.02dB/m，确保足够出光强度。"],
    ["探测器表面 EG", "EG1 = (αBG+αBLB)*EB21", "EG2 = (αBG+αBLB)*EB22", "干涉合束检测。干涉光功率大小决定了接收电信号及散粒噪声的物理限制。"]
]
add_table_with_caption(table1_headers, table1_rows, "集成光学芯片关键传输节点光强和光场状态分布一览表")

make_body("在理想互易性状态下，经历顺时针与逆时针路径相向传输的两束光波在波导合束器表面进行相干叠加。当系统存在由于地球自转或载体旋转引入的输入角速度Ω（rad/s）时，Sagnac效应会导致在两相向波导回路中产生一定的非互易相位差，其计算公式为：")

make_formula(" \\phi_S = \\frac{2\\pi L D}{\\lambda c} \\Omega ", "1.1")

make_body("式中，L为波导环圈的总长度（m），D为硅基集成环圈的几何平均直径（m），λ为发光光源的几何中心波长（nm），c为真空中的光速（2.99792 * 10^8 m/s）。")
make_body("在具体工程中，以波导环长 50m、平均环径 2cm、光源波长 1550nm 为基准参数，可计算得出其光学标度因数为：")

make_formula(" K_c = \\frac{2\\pi L D}{\\lambda c} \\approx 0.01359 rad/(rad/s) ", "1.2")

make_body("这意味着对于 1rad/s 的微小输入角速度，陀螺产生的物理Sagnac相移仅约为 0.014rad。由于该相位差极其微弱，在余弦干涉强度输出曲线上，其静止工作点处于斜率为零的极值点。为解决微弱相移信号的相位敏感度问题并明确旋转方向，必须施加外部偏置相位调制。本课题设计并采用方波和高阶四态方波相位调制。通过在Y分支调制臂上加高频调制电压，使得CW和CCW光波之间在时间延迟τ上形成随时间变化的差分调制相位差：")

make_formula(" \\Delta\\phi_m(t) = \\phi_m(t) - \\phi_m(t-\\tau) ", "1.3")

make_body("式中，τ为光通过集成环圈所需的单程渡越时间（τ = n * L / c，其中n为波导的有效群折射率）。")

add_figure_with_caption("images/report_p14_img1.jpeg", "方波偏置调制波形与探测器干涉强度输出对应关系示意图")

make_body("当输入方波偏置电压且调制频率f0 = 1/(2τ)时，陀螺交替工作在±π/2偏置点上。在没有外部输入角速度（静止状态）时，两个相邻调制半周期内的探测器输出光强完全一致。若存在外部旋转，Sagnac相移φS会导致工作点向相反方向移动，使得相邻两个半周期内的输出光电信号产生明显的幅度差，通过差分提取即可完美解调出旋转速率的幅值和极性。然而，由于片上铌酸锂晶体的温度热膨胀和热光效应，其半波电压Vπ在实际复杂力热环境容易发生严重漂移。这使得传统的开环或简单方波调制会引入极大的温度标度因数误差和死区。")
make_body("为彻底克服由于片上调制器漂移产生的闭环非线性与标度因数漂移问题，本课题联合北京航空航天大学开发了硅光陀螺双闭环数字反馈控制模型。其基本控制逻辑采用四态阶梯方波调制：两个半周期用于解调角速度误差信号，实现Sagnac相移的伺服零跟踪反馈；另两个半周期则用于解调半波电压偏差，实现Vπ的在线自适应反馈。我们利用Simulink工具成功搭建了由光信号干涉模块、时延差分模块、AD量化与FPGA解调模块以及反馈调制信号生成模块组成的动态双闭环闭合控制行为仿真模型。")

add_figure_with_caption("images/report_p16_img1.jpeg", "基于 Simulink 平台的硅光陀螺动态双闭环闭环系统结构设计模型")

make_body("时延差分模块和解调解调控制模块的核心设计及对典型输入的动态闭环仿真响应特性分别如图8、图9、图11和图12所示。仿真时，我们设置AD和DA量化位数为16bit，AD采样周期为 250ns。")

add_figure_with_caption("images/report_p18_img1.jpeg", "（a）时延差分模块输入的连续反馈阶梯调制波形")
add_figure_with_caption("images/report_p18_img3.jpeg", "（b）经环圈渡越时延τ差分对齐后光信号感知的瞬态偏置差分相位波形")
add_figure_with_caption("images/report_p19_img1.jpeg", "（c）FPGA内部数据存储及四个四态积分累加寄存器时序解调逻辑设计")
add_figure_with_caption("images/report_p20_img1.jpeg", "（d）DA前置放大器及波形二次积分阶梯速率斜坡发生器的反馈模型")
add_figure_with_caption("images/report_p21_img1.jpeg", "（e）恒定角速度输入下闭环反馈输出稳定响应（在0.3s内实现零相位精确伺服锁定）")
add_figure_with_caption("images/report_p21_img2.jpeg", "（f）输入突变阶跃信号时闭环测量控制输出的快速暂态阶跃响应特征")

make_body("动态仿真结果清晰表明，该模型在时域不仅能够极为逼真地复现由高频解调与积分引起的暂态超调、AD量化噪声和DA前置放大器热噪声的叠加，还能够精确验证在不同调制频率（如本征频率及其倍数）下，系统标度因数的非互易偏差。它对指导课题5的数字逻辑（FPGA）和ASIC电路（课题4）的控制时序设计具有极强的工程应用实用价值。")

make_heading(3, "微型硅光陀螺基本物理噪声、非互易误差机理与噪声抑制方案", 2)
make_body("为保证微型硅光陀螺的性能极限，我们不仅要建立宏观的电反馈环路，更需要深入物质底层，对光电子和波导材料层面的物理噪声和非互易相位误差进行高精度的微观建模。本课题对此开展了系统化的理论解耦与建模分析，绘制了噪声分类与耦合传递拓扑。")

add_figure_with_caption("images/report_p11_img1.jpeg", "干涉式微型硅光陀螺强度型噪声与非互易相位噪声分类结构流图")

make_body("（1）三类基础物理噪声极限与ARW极限。在硅闭环回路中，探测器提取到的总噪声满足物理上的统计独立和不相关叠加定理，总光电流噪声均方根iN可表示为散粒噪声ish、热噪声ith与RIN强度噪声iRIN的代数平方和：")

make_formula(" i_N^2 = i_{sh}^2 + i_{RIN}^2 + i_{th}^2 = 2 q i_m B_i + i_m^2 \\cdot 10^{\\frac{RIN}{10}} B_i + \\frac{4 k_B T}{R_f} B_i ", "1.4")

make_body("式中，q为电子电量（1.602 * 10^-19 C），im为探测器干涉接收产生的平均光电流，Bi为等效探测带宽（Hz），kB为玻尔兹曼常数，T为开氏温度，Rf为跨阻放大器的反馈跨阻。将该基础流强噪声引入解调公式，可解算出对应的角度随机游走（ARW）极限表达式为：")

make_formula(" ARW = \\frac{2 \\pi}{\\theta_{scale}} \\cdot \\frac{i_N}{\\sqrt{B_i}} \\cdot \\frac{\\lambda c}{2\\pi L D} \\cdot \\frac{180 \\cdot 60}{\\pi} deg/\\sqrt{h} ", "1.5")

make_body("在给定的 50m 硅基集成波导环长和 40mW 入射光功率条件下，散粒噪声和跨阻热相噪构成主要ARW瓶颈。通过仿真发现，当采用 0.75π 最佳调制深度时，陀螺的理论ARW极限达到 0.045°/h^(1/2)；而当进一步通过多态控制把调制深度提高至 0.95π 并提高量子效率η至 0.97 时，由于干涉灵敏度倍增，ARW可显著抑制到 0.0183°/h^(1/2)。")
make_body("（2）热相位噪声（Thermal Phase Noise）。由于光芯片材料的温度振荡和自发热，其由于热光效应引起的折射率局部抖动会在波导回路中引入不可忽视的热阻相位扰动。热相位噪声密度计算公式表示为：")

make_formula(" S_{\\phi, thermal}(f) = \\frac{4 k_B T^2 L \\alpha_L^2 n_{eff}^2}{\\kappa \\pi f} \\cdot [1 - sinc(\\frac{2\\pi f L n_{eff}}{c})] ", "1.6")

make_body("本课题对热相位噪声频谱分布开展了数值计算，发现当将调制工作频率从单倍本征频率调制（1倍，热相位噪声为 9.47 * 10^-17 rad2/Hz）提升至本征频率的奇数高倍频调制时，其周围的热相位噪声发生了极具断崖式的消减：")

add_figure_with_caption("images/report_p24_img1.jpeg", "热相位噪声在不同调制频率（1倍、3倍、5倍本征频率）下的谱密度仿真分布")

make_body("采用 3倍本征频率调制时，热相噪降至 1.05 * 10^-17 rad2/Hz，减小为原来的 11.09%；当采用 5倍本征频率高频偏置时，其相噪极速萎缩到 3.83 * 10^-18 rad2/Hz，仅为原来的 4.04%。这有力论证了“通过高倍频调制，在不牺牲光学长度的前提下强力压制材料自发热热相噪”的物理技术主线，并被集成到后续的芯片驱动（课题4）和仿真中。")
make_body("（3）瑞利背向散射误差（Rayleigh Backscattering）。在超窄硅基集成波导芯层中，由于波导侧壁微观粗糙度和刻蚀条纹不均匀，光波在相向传输中必然存在严重的瑞利反向反射。这些零散散射光与主干涉光发生杂散干涉，会导致顺逆时针光路在解调相差中引入严重的常值相位漂移和快速噪声。")

add_figure_with_caption("images/noise_p7_img1.jpeg", "瑞利背向散射相位耦合物理机制与统计干涉作用模型")

make_body("（4）偏振耦合误差（Polarization Coupling）。硅基光波导即使采用二氧化硅强约束保偏设计，在弯曲和工艺应力下其快慢双轴之间也存在连续的能量串扰，h参数偏大。在光源波导环两端极小的消光比失配下，快轴和慢轴中的传播光在干涉检测时会因光相位不同形成寄生的偏振相位非互易误差。")

add_figure_with_caption("images/noise_p17_img1.jpeg", "硅基集成波导双轴传输中的偏振模式串扰和相关函数相位飘移计算模型")

make_body("（5）克尔非互易效应（Kerr Effect）。在高约束硅基波导芯层（如Si3N4或硅基波导）中，极小的芯层面积使得光功率密度极大（达到 MW/cm^2 量级）。在强光强极化下，材料表现出非线性三阶极化特性，导致介质折射率与光强大小存在正比关系。当Y分支的分束器存在不对称偏置差δ（如49:51）时，两路光的强弱不对称会导致非互易的克尔相移，带来零偏漂移。")

add_figure_with_caption("images/noise_p42_img1.jpeg", "克尔非线性效应模型及不同波导环长度下的相干线宽要求曲线")

make_body("本课题针对这三大非互易误差，建立了精确的数学表达式和抑制映射。我们提出在集成光学芯片中建立对称补偿机制、采用特大谱宽光源（如 35nm 线宽）来破坏散射和偏振光的自相干相干性、并且通过对称性调配对δ进行精细补偿，最终将三大非互易带来的综合零偏不稳定性抑制到 0.2692°/h 级，圆满解决了物理层面的噪声控制瓶颈。")

make_heading(3, "芯片死区（Lock-in Zone）电磁电学耦合物理机理与控制技术", 2)
make_body("在微型高度一体化的硅光陀螺中，光学芯片和前置微弱放大电路的空间间距极窄（通常为 mm 甚至亚微米级）。在高频数字方波调制下，大摆幅的电驱动方波会通过空间极小的分布容抗、电源波动和地回路耦合到探测器（PD）端，使其产生同频的电容性感应串扰相移。这种寄生电容和电阻电磁串扰会导致在小角速度输入时，解调系统因同相电串扰无法实现反馈阶梯波的正常累加，形成致命的测量“死区”（即当角速度低于一定阈值时，陀螺输出保持为零，对旋转毫无反应）。")
make_body("本课题针对该电磁相互作用耦合开展了开创性的芯片级死区物理通路和相移建模。系统分析了包括PCB寄生电容耦合（如-80dB）、电源电压耦合（-90dB）、地平面地回路公共阻抗串扰（-85dB）、外壳静电容性分布耦合（-95dB）及空间高频EMC辐射电磁波耦合（-100dB）五种电学交叉路径。我们成功推导出了决定死区形成边界的总等效误差相位差φe的判据公式。我们得出了如下核心物理结论：由于电串扰相位滞后，当系统满足γ < 1条件时，反馈控制环路会自动将Sagnac有效物理相移湮灭，产生死区锁定。为此，我们设计并开发了死区特征速率一键扫描和寄生通路定量分析算法。")

add_figure_with_caption("images/report_p43_img1.jpeg", "五类芯片级电磁与电学交叉耦合串扰的等效物理电路死区相位误差模型")

# Add cross-talk table
table2_headers = ["电磁串扰物理路径", "典型等效耦合强度", "对死区相位误差贡献", "抑制与电路优化控制方案"]
table2_rows = [
    ["PCB 寄生分布电容耦合", "-80 dB", "高频方波瞬态边沿电平跳变带来 35% 死区贡献", "采用双层屏蔽微带线布线、极化层隔离并增大驱动与PD走线间距。"],
    ["电源公共平面反馈串扰", "-90 dB", "解调电源谐波通过前放放大带来 20% 贡献", "采用芯片级单点LDO供电、LC多级π型去耦滤波，抑制瞬态纹波。"],
    ["公共地回路阻抗交叉耦合", "-85 dB", "瞬态边沿地弹电压抬升引入 25% 贡献", "全差分地信号分布、敏感前放地单点独立接地、严格阻断地回路。"],
    ["外壳及封装容性分布串扰", "-95 dB", "外壳感应电荷高频谐波产生 12% 贡献", "采用导电金属材料进行气密性屏蔽封装，外壳全金属直连数字地。"],
    ["空間高频 EMC 电磁波辐射", "-100 dB", "空中直接耦合天线效应带来 8% 贡献", "关键ASIC和前放电路覆盖专门的屏蔽罩、抑制天线耦合接收谐振。"]
]
add_table_with_caption(table2_headers, table2_rows, "五类关键芯片级电子交叉耦合通路参数、死区贡献及抑制方案")

make_body("基于上述研究，本课题联合中北大学设计了针对死区消除的混合抑制方案。在物理电路设计上，提出基于ASIC全差分电信号结构，使PCB边缘高频电容跳变自动在差分端反相抵消，使耦合衰减压制在-95dB以下；在控制算法上，开发了数字抖动和四态多级斑驳解调策略，通过往Y波导中注入伪随机高频电抖动相位，使系统强制跳出死区不敏感区。通过这些创新方法，使得理论死区半宽被成功压制到近乎零的水平（检测线性度大幅优化），成功破除了妨碍小转速精密测量的死锁，保证了极高的低转速分辨力。")

make_heading(2, "项目取得的阶段性进展及前景", 1)
make_heading(3, "硅光陀螺精密测量模型（课题1）", 2)
make_body("本课题取得的中期阶段性成果展示出了强大的应用前景：通过构建出完整的微型硅光陀螺电光闭环时域测量仿真理论，我国在高度集成的惯性传感领域成功建立起了“物理参数（纳米级波导工艺） -> 链路损耗与噪声拓扑（光子层面） -> 系统标度与伺服控制（电控算法）”的全物理流仿真闭环体系。")
make_body("该精密测量理论在未来具有极高的成果转化效益，它可以直接嵌入到新一代硅光机载导航、无人驾驶微型IMU以及武器制导等微型化角速度传感器设计中，提供低损耗、高稳定性、低死区和宽频带的最佳工艺边界和控制参数组合。它的学术和技术辐射效应，将为全片上集成高稳定光电干涉仪惯性技术的突飞猛进奠定坚实的科学理论基石。")

print("Writing Section III...")
# Section III
make_heading(1, "项目人员及经费投入使用情况", 0)
make_heading(2, "人员及经费投入使用情况", 1)
make_body("根据项目任务书财务管理规范和课题总体预算批复，课题组严格按照项目和财务法规，对专项资金和配套资金的使用进行了精细、合规的管理。")
make_body("（1）课题1“基于硅光技术的高精度角速度测量理论”研究预算批复数总计 544 万元，其中中央财政资金 208 万元，地方或承担单位自筹其他来源资金 336 万元。截至 2026 年 6 月，课题组累计实际到位经费 544 万元（其中中央财政实际到位 208 万元，其他来源配套到位 336 万元），到位率 100%。课题组累计支出资金总计 132.44 万元（其中中央财政资金实际支出 87.12 万元，其他来源资金实际支出 45.32 万元），各项预算科目执行进度良好，资金支出比例与项目研发进展、任务里程碑节点相匹配。")

headers5 = ['科目', '预算批复数 - 中央财政(万元)', '预算批复数 - 其他来源(万元)', '实际到位数 - 中央财政(万元)', '实际到位数 - 其他来源(万元)', '经费支出数 - 中央财政(万元)', '经费支出数 - 其他来源(万元)', '经费执行率 - 中央', '经费执行率 - 其他']
rows5 = [
    ['一、课题支出合计', '208.00', '336.00', '208.00', '336.00', '87.12', '45.32', '41.88%', '13.49%'],
    ['（一）直接费用', '171.70', '336.00', '171.70', '336.00', '68.73', '45.32', '40.03%', '13.49%'],
    ['1. 设备费', '7.00', '10.00', '7.00', '10.00', '0.00', '0.00', '0.00%', '0.00%'],
    ['2. 材料费', '68.20', '176.00', '68.20', '176.00', '22.23', '44.60', '32.59%', '25.34%'],
    ['3. 测试化验加工费', '19.80', '56.00', '19.80', '56.00', '5.61', '0.00', '28.35%', '0.00%'],
    ['4. 会议/差旅/合作费', '21.50', '6.00', '21.50', '6.00', '7.56', '0.72', '35.17%', '11.99%'],
    ['5. 出版/文献/知识产权', '14.00', '10.00', '14.00', '10.00', '5.62', '0.00', '40.16%', '0.00%'],
    ['6. 劳务费', '33.20', '78.00', '33.20', '78.00', '27.01', '0.00', '81.35%', '0.00%'],
    ['7. 专家咨询费', '8.00', '0.00', '8.00', '0.00', '0.70', '0.00', '8.75%', '0.00%'],
    ['（二）间接费用', '36.30', '0.00', '36.30', '0.00', '18.39', '0.00', '50.65%', '0.00%']
]
add_table_with_caption(headers5, rows5, "课题1预算及经费收支执行明细表（与中期情况执行表100%保持一致）")

# Overwrite duplicated other tables by simply omitting them and explaining
make_body("注：本课题1共有三个参与承担单位，分别为浙江大学、北京航空航天大学、中北大学。为提高管理效率并遵循精简不重复的原则，以上经费收支表（表4）已将三家单位的财务数据和主要任务科目进行了系统化的统一汇总，并与《中期情况执行表(1).xlsx》完全保持一致。故不再按单位拆分生成重复冗余的多张表格。")

make_heading(2, "项目经费拨付情况", 1)
make_body("项目牵头单位为北京自动化控制设备研究所。项目牵头单位根据国家任务书经费拨付计划，向本课题1拨付资金：")
make_body("（1）课题1中央财政资金预算总额为 208 万元。课题牵头单位浙江大学累计收到牵头单位拨付的中央财政到账资金 208 万元。浙江大学严格履行合同任务书约定，已向各课题参与单位划拨相关研发资金。浙江大学共收到中央财政到账资金 88 万元；北京航空航天大学共收到到账资金 40 万元；中北大学共收到到账资金 80 万元，划拨执行到位率 100%，极大地保障了各合作单位科研工作的同步推进。")

make_heading(2, "人员及经费调整情况", 1)
make_body("本课题1在中期检查前，各项研究工作、设备测试及人员配置完全在项目和课题任务书的计划大纲轨道内开展。课题组研发团队（教授、博士研究生及科研助理）保持高度稳定。截至2026年6月，课题1中中央财政和地方配套自筹经费的列支全部符合预算规定，未使用任何超出批复科目的支出，未发生任何重大预算科目变更和重大人事调整，财务执行合规。")

print("Writing Section IV...")
# Section IV
make_heading(1, "项目配套支撑条件情况", 0)
make_body("为了完美支撑课题1的高标准科研任务，三个承担单位配备并搭建了国际一流、国内领先的配套硬件设施与实验研发平台：")
make_body("（1）浙江大学配备了硅基集成光电耦合测试平台、超大谱宽掺铒光纤光源测试仪（覆盖 1520nm 至 1610nm 波段）、高精度超窄线宽激光扫频仪、偏振消光比测试仪（消光比测试优于 50dB）、硅片热传导与封装应力控制平台以及极弱信号前置放大示波测试系统，能够完美支撑散粒噪声限制测试、偏振耦合点统计测试、克尔非线性及波导环热相位噪声的系统验证研究。同时组建了高性能GPU理论计算中心，用于开发运行精密测量仿真软件。")
make_body("（2）北京航空航天大学配备了高精度单轴转台、双闭环FPGA高速解调信号测试板（AD采样优于 14bit、采样速率 100MHz）、高速高线性度数模转换反馈链路电路分析仪以及时域振动与力热耦合温箱，用于在真实的物理电路闭环时序中评估和验证本课题建立的双闭环Simulink仿真模型的动态测量误差和参数补偿曲线。")
make_body("（3）中北大学组建了芯片级微纳电磁屏蔽实验室，配备了高阻抗射频地弹测试仪、PCB寄生分布电容电网络解析仪、瞬态电脉冲发生器以及EMC空间辐射耦合天线阵列探头系统，能够精确定量测量和重现PCB走线与集成芯片中的微弱交叉容性串扰强度，为本课题关于芯片死区（Lock-in zone）的物理五路径仿真模型提供精确的实验输入和电参数修正。")

print("Writing Section V...")
# Section V
make_heading(1, "项目组织实施管理工作", 0)
make_heading(2, "项目组织管理情况", 1)
make_body("课题1成立了由浙江大学陈杏藩教授牵头、浙江大学、北京航空航天大学、中北大学核心研究骨干共同参与的课题联合攻关组。自项目启动以来，课题组建立并常态化运行了“月度技术对接、季度专题报告、半年度现场研讨、重大节点联合攻关”的组织管理机制。")
make_body("课题组累计召开了技术专题视频对接会 18 次，围绕“背向散射统计干涉”、“偏振消光测试”及“电串扰寄生参数测量”等重大瓶颈，组织三家单位进行了5次关键阶段的跨地域现场联合实验测试。通过Git仓库和文档协同平台，实现了仿真模型算法、Matlab精密测量代码的统一规范化版本控制。极大地保障了课题内部各单位研究内容和实验参数的高度一致，管理风险降到最低。")

make_heading(2, "项目间协作情况", 1)
make_body("本课题1作为整个项目的“高精度理论基石和仿真模型源头”，与项目内其他课题开展了密切、高效、常态化的参数共享与数据协同。")
make_body("课题1将建立的基本三噪声、瑞利背向散射、偏振耦合及克尔漂移的最优参数边界直接输出给课题2（光敏感芯片）和课题3（光源芯片），用于优化片上硅基集成波导的弯曲半径、h参数（要求 ≤ 1 * 10^-5 m-1）和光源RIN谱密度（要求优于 -126.4 dBc/Hz）。同时，将电磁交叉耦合对死区影响的寄生分布通路和Vπ在线自反馈解调逻辑，传递给课题4（检测ASIC电路）用于指引ASIC前端差分差分走线及FPGA闭环数字抖动模块设计，并为课题5的高稳定微型硅光陀螺整机样机的性能评测提供了完整的仿真平台和测试大纲基准。")

make_body("编制日期：2026年8月4日", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

# ---- Save Document ----
output_filename = os.path.join(BASE_DIR, "1 课题材料模板20260730.docx")
doc.save(output_filename)
print(f"Document saved successfully as '{output_filename}'!")
print(f"File size: {os.path.getsize(output_filename)} bytes")
